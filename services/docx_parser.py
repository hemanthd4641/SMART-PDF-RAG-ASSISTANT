"""
DOCX Parser Service
-------------------
Extracts text from Microsoft Word (.docx) files using python-docx.

Since DOCX files have no fixed page boundaries (unlike PDFs), the entire
document is treated as a single page (page_number=1) to maintain pipeline
compatibility with the chunker, embedder, and retrieval services.

The output format is identical to parse_pdf() and parse_txt(), so no
changes to any downstream service are required.
"""

import os
from typing import Dict, Any
from utils.helpers import get_logger

logger = get_logger("docx_parser")


class DocxParsingError(Exception):
    """Raised when a DOCX file cannot be parsed due to corruption, invalid format, or empty content."""
    pass


def extract_docx_text(file_path: str) -> Dict[str, Any]:
    """
    Parses a DOCX file and extracts all paragraph and table text using python-docx.

    Since DOCX files have no native page concept, the entire document is mapped
    to a single virtual page (page_number=1). This keeps the output format
    fully compatible with parse_pdf() and parse_txt().

    Extraction Strategy:
        1. Extract all paragraph text in document order.
        2. Extract all table cells in row-major order (row by row).
        3. Return combined text as page 1.

    Args:
        file_path: Absolute path to the .docx file.

    Returns:
        Structured output matching the RAG pipeline format:
        {
            "document_name": "sample.docx",
            "pages": [
                {
                    "page_number": 1,
                    "text": "full extracted text content...",
                    "tables": [],
                    "extraction_method": "native"
                }
            ]
        }

    Raises:
        DocxParsingError: If the file is missing, corrupted, empty, or cannot be read.
    """
    document_name = os.path.basename(file_path)

    # Guard: file must exist
    if not os.path.exists(file_path):
        raise DocxParsingError(f"File not found: {file_path}")

    # Guard: file must not be empty
    if os.path.getsize(file_path) == 0:
        raise DocxParsingError(f"DOCX file '{document_name}' is empty (0 bytes).")

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"python-docx failed to open '{document_name}': {e}")
        raise DocxParsingError(
            f"Unable to process document '{document_name}'. "
            f"The file may be corrupted or is not a valid DOCX file."
        )

    text_parts = []

    # 1. Extract paragraph text in document order
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            text_parts.append(stripped)

    # 2. Extract table cell content in row-major order
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                # Join cells with a tab separator to preserve tabular context
                text_parts.append("\t".join(row_cells))

    full_text = "\n".join(text_parts).strip()

    # Guard: document must contain at least some text
    if not full_text:
        raise DocxParsingError(
            f"No text found in document '{document_name}'. "
            f"The file may be empty or contain only images/non-text elements."
        )

    logger.info(
        f"DOCX '{document_name}' parsed successfully: "
        f"{len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, "
        f"{len(full_text)} characters extracted."
    )

    return {
        "document_name": document_name,
        "pages": [
            {
                "page_number": 1,
                "text": full_text,
                "tables": [],          # Tables already merged into text above
                "extraction_method": "native"
            }
        ]
    }
