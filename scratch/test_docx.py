"""
Test cases for DOCX parser service.

Test 1: Valid DOCX with multiple paragraphs  → text extracted successfully
Test 2: Empty DOCX (no paragraphs)           → DocxParsingError raised
Test 3: Corrupted / invalid DOCX file        → DocxParsingError raised
"""
import sys
import os
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from services.docx_parser import extract_docx_text, DocxParsingError

PASS = "[PASS]"
FAIL = "[FAIL]"

results = []

# ─────────────────────────────────────────────
# Test 1: Valid DOCX with multiple paragraphs
# ─────────────────────────────────────────────
print("=" * 60)
print("TEST 1: Valid DOCX with multiple paragraphs")
print("=" * 60)
try:
    from docx import Document

    # Build a real in-memory DOCX with paragraphs and a table
    doc = Document()
    doc.add_heading("ABC Technologies Pvt Ltd", level=1)
    doc.add_paragraph("Founded: 2020")
    doc.add_paragraph("CEO: Sarah Johnson")
    doc.add_paragraph("Headquarters: Bangalore, India")
    doc.add_paragraph("Number of Employees: 500")
    doc.add_paragraph(
        "The company signed a strategic partnership agreement with XYZ Corp in March 2024."
    )

    # Add a simple table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Product"
    table.cell(0, 1).text = "Launch Year"
    table.cell(1, 0).text = "AI Legal Assistant"
    table.cell(1, 1).text = "2022"

    # Save to a temp file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    doc.save(tmp_path)

    result = extract_docx_text(tmp_path)
    os.unlink(tmp_path)

    assert result["document_name"].endswith(".docx"), "document_name must end with .docx"
    assert len(result["pages"]) == 1, "Should have exactly 1 virtual page"
    assert result["pages"][0]["page_number"] == 1, "page_number must be 1"
    assert "ABC Technologies" in result["pages"][0]["text"], "ABC Technologies not found in text"
    assert "Sarah Johnson" in result["pages"][0]["text"], "Sarah Johnson not found in text"
    assert "AI Legal Assistant" in result["pages"][0]["text"], "Table content not found"
    assert result["pages"][0]["extraction_method"] == "native", "extraction_method must be native"

    print(f"{PASS} Text extracted successfully.")
    print(f"       Characters extracted: {len(result['pages'][0]['text'])}")
    print(f"       Page count: {len(result['pages'])}")
    print(f"       extraction_method: {result['pages'][0]['extraction_method']}")
    results.append(True)

except AssertionError as ae:
    print(f"{FAIL} Assertion failed: {ae}")
    results.append(False)
except Exception as e:
    print(f"{FAIL} Unexpected error: {e}")
    results.append(False)

print()

# ─────────────────────────────────────────────
# Test 2: Empty DOCX (no paragraphs)
# ─────────────────────────────────────────────
print("=" * 60)
print("TEST 2: Empty DOCX (no text content)")
print("=" * 60)
try:
    from docx import Document

    # Create a DOCX with no text content
    doc = Document()
    # Add an empty paragraph (stripped → no content)
    doc.add_paragraph("")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    doc.save(tmp_path)

    try:
        extract_docx_text(tmp_path)
        print(f"{FAIL} Expected DocxParsingError was NOT raised.")
        results.append(False)
    except DocxParsingError as e:
        assert "No text found" in str(e), f"Wrong error message: {e}"
        print(f"{PASS} DocxParsingError raised correctly.")
        print(f"       Error message: {e}")
        results.append(True)
    finally:
        os.unlink(tmp_path)

except Exception as e:
    print(f"{FAIL} Unexpected error: {e}")
    results.append(False)

print()

# ─────────────────────────────────────────────
# Test 3: Corrupted / invalid DOCX file
# ─────────────────────────────────────────────
print("=" * 60)
print("TEST 3: Corrupted DOCX file")
print("=" * 60)
try:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, mode="wb") as tmp:
        tmp.write(b"This is not a valid DOCX file. Just random bytes.")
        tmp_path = tmp.name

    try:
        extract_docx_text(tmp_path)
        print(f"{FAIL} Expected DocxParsingError was NOT raised.")
        results.append(False)
    except DocxParsingError as e:
        assert "Unable to process" in str(e) or "corrupted" in str(e).lower(), \
            f"Wrong error message: {e}"
        print(f"{PASS} DocxParsingError raised correctly.")
        print(f"       Error message: {e}")
        results.append(True)
    finally:
        os.unlink(tmp_path)

except Exception as e:
    print(f"{FAIL} Unexpected error: {e}")
    results.append(False)

print()

# ─────────────────────────────────────────────
# Test 4: Pipeline compatibility check
# ─────────────────────────────────────────────
print("=" * 60)
print("TEST 4: Pipeline compatibility (chunker integration)")
print("=" * 60)
try:
    from docx import Document
    from services.chunker import DocumentChunker

    doc = Document()
    for i in range(10):
        doc.add_paragraph(
            f"Paragraph {i+1}: This is a detailed paragraph about the system architecture. "
            f"It describes roles, responsibilities, and key workflows within the application."
        )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    doc.save(tmp_path)

    parsed = extract_docx_text(tmp_path)
    os.unlink(tmp_path)

    chunker = DocumentChunker(chunk_size=500, chunk_overlap=100)
    chunks = chunker.chunk_document(parsed)

    assert len(chunks) > 0, "Chunker produced 0 chunks from DOCX content"
    for chunk in chunks:
        assert "chunk_id" in chunk
        assert "document_name" in chunk
        assert "page_number" in chunk
        assert "chunk_text" in chunk
        assert chunk["page_number"] == 1

    print(f"{PASS} Chunker produced {len(chunks)} chunks from DOCX content.")
    print(f"       All chunks contain required metadata fields.")
    results.append(True)

except AssertionError as ae:
    print(f"{FAIL} Assertion failed: {ae}")
    results.append(False)
except Exception as e:
    print(f"{FAIL} Unexpected error: {e}")
    results.append(False)

print()

# ─────────────────────────────────────────────
# Summary
# ─────────────────────────────────────────────
print("=" * 60)
passed = sum(results)
total = len(results)
print(f"RESULTS: {passed}/{total} tests passed")
if passed == total:
    print("ALL TESTS PASSED")
else:
    failed = [i+1 for i, r in enumerate(results) if not r]
    print(f"FAILED TESTS: {failed}")
print("=" * 60)
